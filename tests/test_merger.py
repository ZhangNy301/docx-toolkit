"""
文档合并模块测试
"""
import os
import pytest
from docx import Document

import sys
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from scripts.core.merger import merge_documents, merge_with_template


class TestMerger:
    """文档合并测试类"""

    def test_merge_documents_empty_list(self, tmp_path):
        """测试空文档列表"""
        output_path = tmp_path / "output.docx"
        with pytest.raises(ValueError, match="文档列表不能为空"):
            merge_documents([], str(output_path))

    def test_merge_documents_invalid_file(self, tmp_path):
        """测试无效文件"""
        # 创建一个非 DOCX 文件
        invalid_file = tmp_path / "invalid.txt"
        invalid_file.write_text("not a docx file")

        # 创建一个有效的 DOCX 文件
        valid_doc = tmp_path / "valid.docx"
        doc = Document()
        doc.add_paragraph("Valid document")
        doc.save(str(valid_doc))

        output_path = tmp_path / "output.docx"
        with pytest.raises(ValueError, match="无效的 DOCX 文件"):
            merge_documents([str(invalid_file), str(valid_doc)], str(output_path))

    def test_merge_documents_single_file(self, tmp_path):
        """测试单个文件合并"""
        # 创建单个文档
        doc_path = tmp_path / "single.docx"
        doc = Document()
        doc.add_paragraph("Single document content")
        doc.save(str(doc_path))

        output_path = tmp_path / "output.docx"
        result = merge_documents([str(doc_path)], str(output_path))

        assert result == str(output_path)
        assert output_path.exists()

        # 验证内容
        result_doc = Document(str(output_path))
        assert "Single document content" in result_doc.paragraphs[0].text

    def test_merge_documents_multiple_files(self, tmp_path):
        """测试多个文件合并"""
        # 创建多个文档
        doc1_path = tmp_path / "doc1.docx"
        doc1 = Document()
        doc1.add_paragraph("Document 1 content")
        doc1.save(str(doc1_path))

        doc2_path = tmp_path / "doc2.docx"
        doc2 = Document()
        doc2.add_paragraph("Document 2 content")
        doc2.save(str(doc2_path))

        doc3_path = tmp_path / "doc3.docx"
        doc3 = Document()
        doc3.add_paragraph("Document 3 content")
        doc3.save(str(doc3_path))

        output_path = tmp_path / "merged.docx"
        result = merge_documents(
            [str(doc1_path), str(doc2_path), str(doc3_path)],
            str(output_path)
        )

        assert result == str(output_path)
        assert output_path.exists()

        # 验证内容
        result_doc = Document(str(output_path))
        paragraphs = [p.text for p in result_doc.paragraphs]
        assert "Document 1 content" in paragraphs
        assert "Document 2 content" in paragraphs
        assert "Document 3 content" in paragraphs

    def test_merge_documents_creates_output_dir(self, tmp_path):
        """测试自动创建输出目录"""
        doc_path = tmp_path / "doc.docx"
        doc = Document()
        doc.add_paragraph("Content")
        doc.save(str(doc_path))

        output_path = tmp_path / "subdir" / "nested" / "output.docx"
        result = merge_documents([str(doc_path)], str(output_path))

        assert result == str(output_path)
        assert output_path.exists()

    def test_merge_with_template(self, tmp_path):
        """测试使用模板合并文档"""
        # 创建模板文档（设置特定的页面尺寸）
        template_path = tmp_path / "template.docx"
        template = Document()
        template.sections[0].page_width = 1224000  # 特定宽度
        template.sections[0].page_height = 1584000  # 特定高度
        template.save(str(template_path))

        # 创建文档
        doc1_path = tmp_path / "doc1.docx"
        doc1 = Document()
        doc1.add_paragraph("Content 1")
        doc1.save(str(doc1_path))

        doc2_path = tmp_path / "doc2.docx"
        doc2 = Document()
        doc2.add_paragraph("Content 2")
        doc2.save(str(doc2_path))

        output_path = tmp_path / "merged_with_template.docx"
        result = merge_with_template(
            [str(doc1_path), str(doc2_path)],
            str(template_path),
            str(output_path)
        )

        assert result == str(output_path)
        assert output_path.exists()

        # 验证页面设置已应用
        # 注意：docxcompose 可能引入微小的舍入差异（< 1mm）
        result_doc = Document(str(output_path))
        assert abs(result_doc.sections[0].page_width - 1224000) < 1000  # 允许 1mm 误差
        assert abs(result_doc.sections[0].page_height - 1584000) < 1000
