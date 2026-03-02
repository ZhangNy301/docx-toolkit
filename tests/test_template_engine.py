"""
模板引擎测试
"""
import os
import pytest
from docx import Document

import sys
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from scripts.core.template_engine import render_template


class TestTemplateEngine:
    """模板引擎测试类"""

    def test_render_simple_variable(self, tmp_path):
        """测试简单变量渲染"""
        # 创建简单模板
        template_path = tmp_path / "template.docx"
        doc = Document()
        doc.add_paragraph("Hello, {{name}}!")
        doc.save(str(template_path))

        # 渲染
        output_path = tmp_path / "output.docx"
        context = {"name": "World"}
        render_template(str(template_path), context, str(output_path))

        # 验证
        result = Document(str(output_path))
        assert "Hello, World!" in result.paragraphs[0].text

    def test_render_with_loop(self, tmp_path):
        """测试循环渲染"""
        template_path = tmp_path / "template.docx"
        doc = Document()
        doc.add_paragraph("{% for item in items %}{{ item }}{% endfor %}")
        doc.save(str(template_path))

        output_path = tmp_path / "output.docx"
        context = {"items": ["A", "B", "C"]}
        render_template(str(template_path), context, str(output_path))

        result = Document(str(output_path))
        assert "ABC" in result.paragraphs[0].text

    def test_render_with_condition(self, tmp_path):
        """测试条件渲染"""
        template_path = tmp_path / "template.docx"
        doc = Document()
        doc.add_paragraph("{% if show %}Visible{% endif %}")
        doc.save(str(template_path))

        output_path = tmp_path / "output.docx"
        context = {"show": True}
        render_template(str(template_path), context, str(output_path))

        result = Document(str(output_path))
        assert "Visible" in result.paragraphs[0].text
