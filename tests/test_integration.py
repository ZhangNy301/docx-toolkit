"""
集成测试
"""
import os
import tempfile
from docx import Document

import sys
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from scripts.core.template_engine import render_template
from scripts.core.merger import merge_documents


class TestIntegration:
    """集成测试类"""

    def test_full_workflow(self):
        """测试完整工作流"""
        with tempfile.TemporaryDirectory() as tmp:
            # 创建模板
            template_path = os.path.join(tmp, "template.docx")
            doc = Document()
            doc.add_paragraph("Hello, {{name}}!")
            doc.save(template_path)

            # 渲染
            output_path = os.path.join(tmp, "output.docx")
            render_template(template_path, {"name": "World"}, output_path)

            # 验证
            result = Document(output_path)
            assert "Hello, World!" in result.paragraphs[0].text

    def test_template_and_merge_workflow(self):
        """测试模板渲染后合并"""
        with tempfile.TemporaryDirectory() as tmp:
            # 创建模板
            template_path = os.path.join(tmp, "template.docx")
            doc = Document()
            doc.add_paragraph("Chapter: {{title}}")
            doc.save(template_path)

            # 渲染两个文档
            doc1_path = os.path.join(tmp, "doc1.docx")
            doc2_path = os.path.join(tmp, "doc2.docx")
            render_template(template_path, {"title": "Introduction"}, doc1_path)
            render_template(template_path, {"title": "Methods"}, doc2_path)

            # 合并
            merged_path = os.path.join(tmp, "merged.docx")
            merge_documents([doc1_path, doc2_path], merged_path)

            # 验证
            result = Document(merged_path)
            texts = [p.text for p in result.paragraphs]
            assert any("Introduction" in t for t in texts)
            assert any("Methods" in t for t in texts)
