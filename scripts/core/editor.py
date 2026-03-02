"""
DOCX 精细编辑模块

提供基于 unpack-edit-pack 模式的精细编辑功能
"""
import os
import re
import shutil
import tempfile
import uuid
from pathlib import Path
from typing import Dict, Any, Optional, List, Tuple
from xml.etree import ElementTree as ET

from docx import Document

from .utils import unpack_docx, pack_docx, get_next_image_name, validate_docx


# XML 命名空间
NAMESPACES = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
    'v': 'urn:schemas-microsoft-com:vml',
    'o': 'urn:schemas-microsoft-com:office:office',
}

# 注册命名空间
for prefix, uri in NAMESPACES.items():
    ET.register_namespace(prefix, uri)


def _get_content_type(extension: str) -> str:
    """根据文件扩展名获取 Content-Type"""
    content_types = {
        '.png': 'image/png',
        '.jpg': 'image/jpeg',
        '.jpeg': 'image/jpeg',
        '.gif': 'image/gif',
        '.bmp': 'image/bmp',
        '.tiff': 'image/tiff',
        '.tif': 'image/tiff',
        '.emf': 'image/x-emf',
        '.wmf': 'image/x-wmf',
    }
    return content_types.get(extension.lower(), 'image/png')


def _mm_to_emu(mm: float) -> int:
    """毫米转换为 EMU (English Metric Unit)"""
    # 1 inch = 25.4 mm = 914400 EMU
    return int(mm * 914400 / 25.4)


def _generate_rid(relationships: dict) -> str:
    """生成新的 rId"""
    max_id = 0
    for rid in relationships.keys():
        if rid.startswith('rId'):
            try:
                num = int(rid[3:])
                max_id = max(max_id, num)
            except ValueError:
                pass
    return f'rId{max_id + 1}'


def _parse_relationships(rels_path: str) -> Tuple[dict, ET.ElementTree]:
    """解析 relationships 文件"""
    tree = ET.parse(rels_path)
    root = tree.getroot()

    relationships = {}
    for rel in root.findall('.//{http://schemas.openxmlformats.org/package/2006/relationships}Relationship'):
        rid = rel.get('Id')
        target = rel.get('Target')
        rtype = rel.get('Type')
        relationships[rid] = {
            'target': target,
            'type': rtype
        }

    return relationships, tree


def _add_relationship(rels_path: str, target: str, rel_type: str) -> str:
    """添加新的 relationship 并返回 rId"""
    relationships, tree = _parse_relationships(rels_path)
    rid = _generate_rid(relationships)

    root = tree.getroot()
    ns = '{http://schemas.openxmlformats.org/package/2006/relationships}'

    new_rel = ET.SubElement(root, f'{ns}Relationship')
    new_rel.set('Id', rid)
    new_rel.set('Type', rel_type)
    new_rel.set('Target', target)

    tree.write(rels_path, xml_declaration=True, encoding='UTF-8')
    return rid


def _update_content_types(content_types_path: str, extension: str, content_type: str):
    """更新 [Content_Types].xml 添加新的扩展名类型"""
    tree = ET.parse(content_types_path)
    root = tree.getroot()

    ns = '{http://schemas.openxmlformats.org/package/2006/content-types}'

    # 检查是否已存在
    for override in root.findall(f'{ns}Default'):
        if override.get('Extension') == extension:
            return

    # 添加新的默认类型
    new_default = ET.SubElement(root, f'{ns}Default')
    new_default.set('Extension', extension)
    new_default.set('ContentType', content_type)

    tree.write(content_types_path, xml_declaration=True, encoding='UTF-8')


def _create_drawing_xml(
    rid: str,
    width_emu: int,
    height_emu: int,
    name: str = "Picture",
    descr: str = ""
) -> str:
    """创建图片的 drawing XML"""
    drawing_xml = f'''<w:drawing xmlns:w="{NAMESPACES['w']}" xmlns:r="{NAMESPACES['r']}" xmlns:a="{NAMESPACES['a']}" xmlns:pic="{NAMESPACES['pic']}" xmlns:wp="{NAMESPACES['wp']}">
  <wp:inline distT="0" distB="0" distL="0" distR="0">
    <wp:extent cx="{width_emu}" cy="{height_emu}"/>
    <wp:docPr id="{abs(hash(name) % 1000000)}" name="{name}" descr="{descr}"/>
    <wp:cNvGraphicFramePr>
      <a:graphicFrameLocks xmlns:a="{NAMESPACES['a']}" noChangeAspect="1"/>
    </wp:cNvGraphicFramePr>
    <a:graphic xmlns:a="{NAMESPACES['a']}">
      <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
        <pic:pic xmlns:pic="{NAMESPACES['pic']}">
          <pic:nvPicPr>
            <pic:cNvPr id="0" name="{name}"/>
            <pic:cNvPicPr/>
          </pic:nvPicPr>
          <pic:blipFill>
            <a:blip xmlns:r="{NAMESPACES['r']}" r:embed="{rid}"/>
            <a:stretch>
              <a:fillRect/>
            </a:stretch>
          </pic:blipFill>
          <pic:spPr>
            <a:xfrm>
              <a:off x="0" y="0"/>
              <a:ext cx="{width_emu}" cy="{height_emu}"/>
            </a:xfrm>
            <a:prstGeom prst="rect">
              <a:avLst/>
            </a:prstGeom>
          </pic:spPr>
        </pic:pic>
      </a:graphicData>
    </a:graphic>
  </wp:inline>
</w:drawing>'''
    return drawing_xml


def _get_image_dimensions(image_path: str) -> Tuple[int, int]:
    """获取图片尺寸 (像素)"""
    try:
        from PIL import Image
        with Image.open(image_path) as img:
            return img.size  # (width, height)
    except ImportError:
        # 如果没有 PIL，返回默认尺寸
        return 300, 200
    except Exception:
        return 300, 200


def _calculate_dimensions(
    image_path: str,
    width_mm: float,
    height_mm: Optional[float] = None,
    dpi: int = 96
) -> Tuple[int, int]:
    """计算图片的 EMU 尺寸"""
    if height_mm is not None:
        return _mm_to_emu(width_mm), _mm_to_emu(height_mm)

    # 根据原始宽高比计算高度
    orig_width, orig_height = _get_image_dimensions(image_path)

    if orig_width > 0:
        aspect_ratio = orig_height / orig_width
        height_mm = width_mm * aspect_ratio
    else:
        height_mm = width_mm * 0.75  # 默认 4:3 比例

    return _mm_to_emu(width_mm), _mm_to_emu(height_mm)


def insert_image(
    docx_path: str,
    image_path: str,
    output_path: str,
    position: Optional[Dict[str, Any]] = None,
    caption: Optional[str] = None,
    width_mm: float = 150,
    height_mm: Optional[float] = None,
    paragraph_index: Optional[int] = None,
    insert_after_marker: Optional[str] = None
) -> str:
    """
    在文档中插入图片

    使用 unpack-edit-pack 模式实现精细控制

    Args:
        docx_path: 源 DOCX 文件路径
        image_path: 要插入的图片路径
        output_path: 输出文件路径
        position: 位置配置 (保留用于扩展)
        caption: 图片标题 (可选)
        width_mm: 图片宽度 (毫米), 默认 150mm
        height_mm: 图片高度 (毫米), 默认自动计算
        paragraph_index: 插入到指定段落后 (0-indexed), None 表示追加到末尾
        insert_after_marker: 在包含此文本的段落后插入

    Returns:
        输出文件路径

    Raises:
        FileNotFoundError: 源文件或图片不存在
        ValueError: 无效的 DOCX 文件
    """
    # 验证输入
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"DOCX 文件不存在: {docx_path}")
    if not os.path.exists(image_path):
        raise FileNotFoundError(f"图片文件不存在: {image_path}")
    if not validate_docx(docx_path):
        raise ValueError(f"无效的 DOCX 文件: {docx_path}")

    # 确保输出目录存在
    output_dir = os.path.dirname(output_path)
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)

    # Step 1: 解压 DOCX
    temp_dir = unpack_docx(docx_path)

    try:
        # Step 2: 复制图片到 media 目录
        media_dir = os.path.join(temp_dir, 'word', 'media')
        if not os.path.exists(media_dir):
            os.makedirs(media_dir)

        # 获取新图片名称
        new_image_name, _ = get_next_image_name(temp_dir)
        image_ext = os.path.splitext(image_path)[1]
        if image_ext:
            new_image_name = os.path.splitext(new_image_name)[0] + image_ext

        target_image_path = os.path.join(media_dir, new_image_name)
        shutil.copy2(image_path, target_image_path)

        # Step 3: 更新 [Content_Types].xml
        content_types_path = os.path.join(temp_dir, '[Content_Types].xml')
        if image_ext:
            content_type = _get_content_type(image_ext)
            _update_content_types(content_types_path, image_ext.lstrip('.'), content_type)

        # Step 4: 更新 relationships
        rels_dir = os.path.join(temp_dir, 'word', '_rels')
        if not os.path.exists(rels_dir):
            os.makedirs(rels_dir)

        document_rels_path = os.path.join(rels_dir, 'document.xml.rels')

        # 如果 relationships 文件不存在，创建一个
        if not os.path.exists(document_rels_path):
            rels_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
</Relationships>'''
            with open(document_rels_path, 'w', encoding='utf-8') as f:
                f.write(rels_xml)

        # 添加图片关系
        image_rel_type = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image'
        rid = _add_relationship(document_rels_path, f'media/{new_image_name}', image_rel_type)

        # Step 5: 插入图片 XML 到 document.xml
        document_path = os.path.join(temp_dir, 'word', 'document.xml')

        # 计算图片尺寸
        width_emu, height_emu = _calculate_dimensions(image_path, width_mm, height_mm)

        # 创建 drawing XML
        drawing_xml = _create_drawing_xml(
            rid=rid,
            width_emu=width_emu,
            height_emu=height_emu,
            name=new_image_name,
            descr=caption or ""
        )

        # 解析 document.xml
        doc_tree = ET.parse(document_path)
        doc_root = doc_tree.getroot()

        # 查找 body
        body = doc_root.find('.//w:body', NAMESPACES)
        if body is None:
            body = doc_root

        # 创建包含图片的段落
        pic_para = ET.fromstring(f'''<w:p xmlns:w="{NAMESPACES['w']}">
  <w:pPr>
    <w:jc w:val="center"/>
  </w:pPr>
  <w:r>
    {drawing_xml}
  </w:r>
</w:p>''')

        # 如果有标题，添加标题段落
        if caption:
            caption_para = ET.fromstring(f'''<w:p xmlns:w="{NAMESPACES['w']}">
  <w:pPr>
    <w:jc w:val="center"/>
    <w:pStyle w:val="Caption"/>
  </w:pPr>
  <w:r>
    <w:t>{caption}</w:t>
  </w:r>
</w:p>''')
        else:
            caption_para = None

        # 确定插入位置
        paragraphs = body.findall('.//w:p', NAMESPACES)

        insert_index = len(paragraphs)  # 默认追加到末尾

        if insert_after_marker is not None:
            # 查找包含标记的段落
            for i, para in enumerate(paragraphs):
                text_elements = para.findall('.//w:t', NAMESPACES)
                for t in text_elements:
                    if t.text and insert_after_marker in t.text:
                        insert_index = i + 1
                        break
                else:
                    continue
                break
        elif paragraph_index is not None:
            insert_index = min(paragraph_index + 1, len(paragraphs))

        # 插入图片段落
        # 注意：sectPr 必须是 body 的最后一个子元素，所以要在它之前插入
        sectPr = body.find('.//w:sectPr', NAMESPACES)

        if insert_index >= len(paragraphs):
            # 追加到末尾，但在 sectPr 之前
            if sectPr is not None:
                sect_index = list(body).index(sectPr)
                body.insert(sect_index, pic_para)
                if caption_para is not None:
                    body.insert(sect_index + 1, caption_para)
            else:
                body.append(pic_para)
                if caption_para is not None:
                    body.append(caption_para)
        else:
            target_para = paragraphs[insert_index]
            target_index = list(body).index(target_para)
            body.insert(target_index, pic_para)
            if caption_para is not None:
                body.insert(target_index + 1, caption_para)

        # 保存修改后的 document.xml
        doc_tree.write(document_path, xml_declaration=True, encoding='UTF-8')

        # Step 6: 打包回 DOCX
        pack_docx(temp_dir, output_path)

    finally:
        # 清理临时目录
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)

    return output_path


def find_and_replace(docx_path: str, replacements: Dict[str, str], output_path: str) -> str:
    """
    查找并替换文本

    Args:
        docx_path: 源 DOCX 文件路径
        replacements: 替换字典 {旧文本: 新文本}
        output_path: 输出文件路径

    Returns:
        输出文件路径
    """
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"DOCX 文件不存在: {docx_path}")
    if not validate_docx(docx_path):
        raise ValueError(f"无效的 DOCX 文件: {docx_path}")

    # 确保输出目录存在
    output_dir = os.path.dirname(output_path)
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)

    doc = Document(docx_path)

    # 替换段落中的文本
    for para in doc.paragraphs:
        for old, new in replacements.items():
            for run in para.runs:
                if old in run.text:
                    run.text = run.text.replace(old, new)

    # 替换表格中的文本
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for old, new in replacements.items():
                        for run in para.runs:
                            if old in run.text:
                                run.text = run.text.replace(old, new)

    # 替换页眉页脚中的文本
    for section in doc.sections:
        # 页眉
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header is not None:
                for para in header.paragraphs:
                    for old, new in replacements.items():
                        for run in para.runs:
                            if old in run.text:
                                run.text = run.text.replace(old, new)
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                for old, new in replacements.items():
                                    for run in para.runs:
                                        if old in run.text:
                                            run.text = run.text.replace(old, new)

        # 页脚
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer is not None:
                for para in footer.paragraphs:
                    for old, new in replacements.items():
                        for run in para.runs:
                            if old in run.text:
                                run.text = run.text.replace(old, new)
                for table in footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                for old, new in replacements.items():
                                    for run in para.runs:
                                        if old in run.text:
                                            run.text = run.text.replace(old, new)

    doc.save(output_path)
    return output_path


def insert_paragraph(
    docx_path: str,
    text: str,
    output_path: str,
    paragraph_index: Optional[int] = None,
    insert_after_marker: Optional[str] = None,
    style: Optional[str] = None
) -> str:
    """
    在文档中插入段落

    Args:
        docx_path: 源 DOCX 文件路径
        text: 段落文本
        output_path: 输出文件路径
        paragraph_index: 插入到指定段落后 (0-indexed), None 表示追加到末尾
        insert_after_marker: 在包含此文本的段落后插入
        style: 段落样式名称

    Returns:
        输出文件路径
    """
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"DOCX 文件不存在: {docx_path}")
    if not validate_docx(docx_path):
        raise ValueError(f"无效的 DOCX 文件: {docx_path}")

    # 确保输出目录存在
    output_dir = os.path.dirname(output_path)
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)

    doc = Document(docx_path)

    # 确定插入位置
    insert_index = len(doc.paragraphs)

    if insert_after_marker is not None:
        for i, para in enumerate(doc.paragraphs):
            if insert_after_marker in para.text:
                insert_index = i + 1
                break
    elif paragraph_index is not None:
        insert_index = min(paragraph_index + 1, len(doc.paragraphs))

    # 创建新段落
    if insert_index >= len(doc.paragraphs):
        new_para = doc.add_paragraph(text, style=style)
    else:
        # 在指定位置插入
        new_para = doc.paragraphs[insert_index].insert_paragraph_before(text)
        if style:
            new_para.style = style

    doc.save(output_path)
    return output_path


def delete_paragraph(
    docx_path: str,
    output_path: str,
    paragraph_index: Optional[int] = None,
    contains_text: Optional[str] = None
) -> str:
    """
    删除文档中的段落

    Args:
        docx_path: 源 DOCX 文件路径
        output_path: 输出文件路径
        paragraph_index: 要删除的段落索引 (0-indexed)
        contains_text: 删除包含此文本的段落

    Returns:
        输出文件路径
    """
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"DOCX 文件不存在: {docx_path}")
    if not validate_docx(docx_path):
        raise ValueError(f"无效的 DOCX 文件: {docx_path}")

    # 确保输出目录存在
    output_dir = os.path.dirname(output_path)
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)

    doc = Document(docx_path)

    if contains_text is not None:
        # 删除包含特定文本的段落
        paragraphs_to_delete = []
        for i, para in enumerate(doc.paragraphs):
            if contains_text in para.text:
                paragraphs_to_delete.append(para)

        for para in paragraphs_to_delete:
            p = para._element
            p.getparent().remove(p)

    elif paragraph_index is not None and 0 <= paragraph_index < len(doc.paragraphs):
        para = doc.paragraphs[paragraph_index]
        p = para._element
        p.getparent().remove(p)

    doc.save(output_path)
    return output_path


def get_document_text(docx_path: str) -> str:
    """
    获取文档中的所有文本

    Args:
        docx_path: DOCX 文件路径

    Returns:
        文档文本内容
    """
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"DOCX 文件不存在: {docx_path}")
    if not validate_docx(docx_path):
        raise ValueError(f"无效的 DOCX 文件: {docx_path}")

    doc = Document(docx_path)
    text_parts = []

    for para in doc.paragraphs:
        text_parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text_parts.append(para.text)

    return '\n'.join(text_parts)


def find_text(docx_path: str, search_text: str) -> List[Dict[str, Any]]:
    """
    在文档中查找文本

    Args:
        docx_path: DOCX 文件路径
        search_text: 要查找的文本

    Returns:
        匹配结果列表 [{'paragraph_index': int, 'text': str, 'context': str}, ...]
    """
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"DOCX 文件不存在: {docx_path}")
    if not validate_docx(docx_path):
        raise ValueError(f"无效的 DOCX 文件: {docx_path}")

    doc = Document(docx_path)
    results = []

    for i, para in enumerate(doc.paragraphs):
        if search_text in para.text:
            results.append({
                'paragraph_index': i,
                'text': para.text,
                'context': para.text[:100] + '...' if len(para.text) > 100 else para.text,
                'location': 'body'
            })

    # 也搜索表格
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    if search_text in para.text:
                        results.append({
                            'table_index': table_idx,
                            'row_index': row_idx,
                            'cell_index': cell_idx,
                            'text': para.text,
                            'location': 'table'
                        })

    return results
