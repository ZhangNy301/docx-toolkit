"""
DOCX 文档合并模块
"""
import os
from typing import List

from docx import Document
from docxcompose.composer import Composer

from .utils import validate_docx


def merge_documents(
    doc_paths: List[str],
    output_path: str,
    page_break: bool = True,
    preserve_format: bool = True
) -> str:
    """合并多个 DOCX 文档"""
    if not doc_paths:
        raise ValueError("文档列表不能为空")

    for path in doc_paths:
        if not validate_docx(path):
            raise ValueError(f"无效的 DOCX 文件: {path}")

    output_dir = os.path.dirname(output_path)
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)

    master_doc = Document(doc_paths[0])
    composer = Composer(master_doc)

    for path in doc_paths[1:]:
        doc = Document(path)
        composer.append(doc)

    composer.save(output_path)
    return output_path


def merge_with_template(
    doc_paths: List[str],
    template_path: str,
    output_path: str
) -> str:
    """使用模板合并文档（应用统一的页面设置）"""
    template = Document(template_path)
    template_section = template.sections[0]

    master_doc = Document(doc_paths[0])

    for section in master_doc.sections:
        section.page_width = template_section.page_width
        section.page_height = template_section.page_height
        section.top_margin = template_section.top_margin
        section.bottom_margin = template_section.bottom_margin
        section.left_margin = template_section.left_margin
        section.right_margin = template_section.right_margin

    composer = Composer(master_doc)

    for path in doc_paths[1:]:
        doc = Document(path)
        composer.append(doc)

    composer.save(output_path)
    return output_path
