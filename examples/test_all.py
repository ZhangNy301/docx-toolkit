#!/usr/bin/env python3
"""
docx-toolkit 功能测试脚本

测试三个核心功能:
1. 模板渲染 (template_engine)
2. 文档合并 (merger)
3. 图片插入 (editor)
"""
import os
import sys

# 添加父目录到路径
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from scripts.core.template_engine import render_template
from scripts.core.merger import merge_documents
from scripts.core.editor import insert_image

def test_template_render():
    """测试1: 模板渲染"""
    print("\n" + "="*50)
    print("测试1: 模板渲染")
    print("="*50)

    template_path = "template.docx"
    data_path = "data.json"
    output_path = "output_rendered.docx"

    import json
    with open(data_path, 'r', encoding='utf-8') as f:
        context = json.load(f)

    result = render_template(template_path, context, output_path)
    print(f"✓ 模板渲染完成: {result}")

    # 验证
    from docx import Document
    doc = Document(result)
    print(f"  文档包含 {len(doc.paragraphs)} 个段落")
    print(f"  第一个段落: {doc.paragraphs[0].text[:30]}...")
    return result

def test_merge():
    """测试2: 文档合并"""
    print("\n" + "="*50)
    print("测试2: 文档合并")
    print("="*50)

    doc_paths = [
        "chapter1.docx",
        "chapter2.docx",
        "chapter3.docx"
    ]
    output_path = "output_merged.docx"

    result = merge_documents(doc_paths, output_path)
    print(f"✓ 文档合并完成: {result}")

    # 验证
    from docx import Document
    doc = Document(result)
    print(f"  合并后文档包含 {len(doc.paragraphs)} 个段落")
    for i, p in enumerate(doc.paragraphs[:6]):
        if p.text.strip():
            print(f"  段落{i+1}: {p.text}")
    return result

def test_insert_image():
    """测试3: 图片插入"""
    print("\n" + "="*50)
    print("测试3: 图片插入")
    print("="*50)

    doc_path = "output_rendered.docx"
    image_path = "figure.png"
    output_path = "output_with_image.docx"

    # 在 "摘要" 后面插入图片
    position = {"type": "after_text", "value": "摘要"}
    caption = "图1. 测试图片示例"

    result = insert_image(doc_path, image_path, output_path, position, caption)
    print(f"✓ 图片插入完成: {result}")

    # 验证
    from docx import Document
    doc = Document(result)
    print(f"  文档包含 {len(doc.paragraphs)} 个段落")

    # 检查图片
    from scripts.core.utils import get_media_files
    media = get_media_files(output_path)
    print(f"  文档包含 {len(media)} 个媒体文件")
    return result

def main():
    """运行所有测试"""
    print("\n" + "="*60)
    print("       docx-toolkit 功能测试")
    print("="*60)

    # 切换到 examples 目录
    script_dir = os.path.dirname(os.path.abspath(__file__))
    os.chdir(script_dir)
    print(f"工作目录: {os.getcwd()}")

    try:
        # 测试1: 模板渲染
        test_template_render()

        # 测试2: 文档合并
        test_merge()

        # 测试3: 图片插入
        test_insert_image()

        print("\n" + "="*60)
        print("       ✓ 所有测试完成！")
        print("="*60)
        print("\n生成的文件:")
        for f in ["output_rendered.docx", "output_merged.docx", "output_with_image.docx"]:
            if os.path.exists(f):
                size = os.path.getsize(f)
                print(f"  • {f} ({size:,} bytes)")

        print("\n你可以用 Word 或 WPS 打开这些文件查看效果。")

    except Exception as e:
        print(f"\n✗ 测试失败: {e}")
        import traceback
        traceback.print_exc()
        return 1

    return 0

if __name__ == "__main__":
    sys.exit(main())
